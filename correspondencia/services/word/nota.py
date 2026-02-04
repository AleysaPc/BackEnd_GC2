from io import BytesIO
from django.utils.html import strip_tags
import html

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


from ..utils.fechas import formatear_fecha_es
from ..word.base import crear_documento, obtener_fecha
from ..utils.formato import sin_espacios


def generar_nota_word(correspondenciaElaborada):
    doc = crear_documento()

    fecha = obtener_fecha(correspondenciaElaborada.fecha_envio)
    fecha_envio_str = formatear_fecha_es(fecha)

    p_fecha = doc.add_paragraph(f"La Paz, {fecha_envio_str}")
    sin_espacios(p_fecha)

    parrafo_cite = doc.add_paragraph()
    run_cite = parrafo_cite.add_run(correspondenciaElaborada.cite)
    run_cite.bold = True

    p_senor = doc.add_paragraph("Señor:")
    sin_espacios(p_senor)

    contacto = correspondenciaElaborada.contacto
    destino_interno = getattr(correspondenciaElaborada, "destino_interno", None)

    if correspondenciaElaborada.ambito == "interno" and destino_interno:
        nombre = " ".join(
            [
                destino_interno.first_name or "",
                destino_interno.second_name or "",
                destino_interno.last_name or "",
                destino_interno.second_last_name or "",
            ]
        ).strip()

        if nombre:
            p_nombre = doc.add_paragraph(nombre)
            sin_espacios(p_nombre)

        if destino_interno.cargo:
            p_cargo = doc.add_paragraph()
            r = p_cargo.add_run(destino_interno.cargo.upper())
            r.bold = True
            sin_espacios(p_cargo)

        if destino_interno.departamento:
            p_depto = doc.add_paragraph()
            r = p_depto.add_run(str(destino_interno.departamento).upper())
            r.bold = True
            sin_espacios(p_depto)

    elif contacto:
        titulo_dict = {
            "Ingeniero": "Ing.",
            "Licenciado": "Lic.",
            "Doctor": "Dr.",
            "Abogado": "Abog.",
            "Profesor": "Prof.",
            "Magister": "Mgs.",
        }

        titulo = titulo_dict.get(contacto.titulo_profesional, "")
        nombre = f"{titulo} {contacto.nombre_contacto or ''} {contacto.apellido_pat_contacto or ''} {contacto.apellido_mat_contacto or ''}".strip()

        p_nombre = doc.add_paragraph(nombre)
        sin_espacios(p_nombre)

        if contacto.cargo:
            p_cargo = doc.add_paragraph()
            r = p_cargo.add_run(contacto.cargo.upper())
            r.bold = True
            sin_espacios(p_cargo)

        if contacto.institucion:
            p_inst = doc.add_paragraph()
            r = p_inst.add_run(str(contacto.institucion).upper())
            r.bold = True
            sin_espacios(p_inst)

    p_presente = doc.add_paragraph("Presente.-")
    sin_espacios(p_presente)

    parrafo_ref = doc.add_paragraph()
    parrafo_ref.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_ref = parrafo_ref.add_run(f"Ref.: {correspondenciaElaborada.referencia}")
    run_ref.bold = True
    run_ref.underline = True

    doc.add_paragraph("De nuestra mayor consideración:")

    texto = strip_tags(correspondenciaElaborada.descripcion_desarrollo)
    texto = html.unescape(texto)

    p_desarrollo = doc.add_paragraph(texto)
    p_desarrollo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph(
        "Sin otro particular, nos despedimos con las consideraciones más distinguidas."
    )

    p_atentamente = doc.add_paragraph()
    r = p_atentamente.add_run("Atentamente,")
    r.bold = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"correspondencia_{correspondenciaElaborada.cite}.docx"
    return buffer, filename

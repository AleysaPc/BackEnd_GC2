from io import BytesIO
from django.utils.html import strip_tags
import html
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from ..utils.fechas import formatear_fecha_es
from ..utils.formato import sin_espacios, agregar_linea_divisora
from ..word.base import obtener_fecha

def generar_informe_word(correspondenciaElaborada):
    """
    Genera un documento Word con formato de Informe.
    Basado en tu HTML: encabezado con De:, Para:, Ref:, Fecha: y secciones Introducción, Desarrollo, Conclusión
    """

    doc = Document()

    # 🔹 TITULO
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("INFORME")
    run_titulo.bold = True
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_titulo.bold = True
    p_titulo.runs[0].font.size = Pt(16)
    sin_espacios(p_titulo)


    # 🔹 CITE
    p_cite = doc.add_paragraph(correspondenciaElaborada.cite)
    p_cite.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_cite.runs[0].font.size = Pt(14)
    

    # 🔹 DE
    usuario_obj = getattr(correspondenciaElaborada, "usuario", None)
    if usuario_obj:
        usuario = " ".join(
            filter(
                None,
                [
                    getattr(usuario_obj, "first_name", ""),
                    getattr(usuario_obj, "second_name", ""),
                    getattr(usuario_obj, "last_name", ""),
                    getattr(usuario_obj, "second_last_name", ""),
                ],
            )
        ).strip() or str(usuario_obj)
    else:
        usuario = "Desconocido"
    p_de = doc.add_paragraph()
    run_label = p_de.add_run("De: ")
    run_label.bold = True
    run_text = p_de.add_run(usuario)
    run_text.bold = False
    sin_espacios(p_de)

    # 🔹 PARA
    contacto = getattr(correspondenciaElaborada, "contacto", None)
    destino_interno = getattr(correspondenciaElaborada, "destino_interno", None)
    if correspondenciaElaborada.ambito == "interno" and destino_interno:
        nombre_completo = " ".join(
            [
                destino_interno.first_name or "",
                destino_interno.second_name or "",
                destino_interno.last_name or "",
                destino_interno.second_last_name or "",
            ]
        ).strip()
    elif contacto:
        titulo_dict = {
            "Ingeniero": "Ing.",
            "Licenciado": "Lic.",
            "Doctor": "Dr.",
            "Abogado": "Abog.",
            "Profesor": "Prof.",
            "Magister": "Mgs.",
        }
        titulo = titulo_dict.get(getattr(contacto, "titulo_profesional", None), "")
        nombre_completo = f"{titulo} {contacto.nombre_contacto or ''} {contacto.apellido_pat_contacto or ''} {contacto.apellido_mat_contacto or ''}".strip()
    else:
        nombre_completo = ""

    if nombre_completo:
        p_para = doc.add_paragraph()
        run_label = p_para.add_run("Para: ")
        run_label.bold = True
        run_text = p_para.add_run(nombre_completo)
        run_text.bold = False
        sin_espacios(p_para)

    # 🔹 REF
    p_ref = doc.add_paragraph()
    run_label = p_ref.add_run("Ref.: ")
    run_label.bold = True
    p_ref.add_run(str(correspondenciaElaborada.referencia))
    sin_espacios(p_ref)

    #FECHA

    fecha = obtener_fecha(correspondenciaElaborada.fecha_envio)
    fecha_envio_str = formatear_fecha_es(fecha)

    p_fecha = doc.add_paragraph()
    run_label = p_fecha.add_run("Fecha: ")
    run_label.bold = True
    p_fecha.add_run(str(fecha_envio_str))
    agregar_linea_divisora(p_fecha)


    doc.add_paragraph()  # espacio antes del cuerpo

    # 🔹 CUERPO: Introducción, Desarrollo, Conclusión
    if hasattr(correspondenciaElaborada, "descripcion_introduccion"):
        intro = strip_tags(correspondenciaElaborada.descripcion_introduccion or "")
        
        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        run_label = p_intro.add_run("Introducción:\n")
        run_label.bold = True

        p_intro.add_run(intro)

    if hasattr(correspondenciaElaborada, "descripcion_desarrollo"):
        desarrollo = strip_tags(correspondenciaElaborada.descripcion_desarrollo or "")
        
        p_desarrollo  = doc.add_paragraph()
        p_desarrollo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        run_label = p_desarrollo.add_run("Desarrollo:\n")
        run_label.bold = True

        p_desarrollo.add_run(desarrollo)

    if hasattr(correspondenciaElaborada, "descripcion_conclusion"):
        conclusion = strip_tags(correspondenciaElaborada.descripcion_conclusion or "")
        
        p_conclusion = doc.add_paragraph()
        p_conclusion.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        run_label = p_conclusion.add_run("Conclusion:\n")
        run_label.bold = True

        p_conclusion.add_run(conclusion)
    doc.add_paragraph()  # espacio antes de cierre

    # 🔹 SALUDO DE CIERRE
    p_cierre = doc.add_paragraph()
    run_cierre = p_cierre.add_run("Atentamente,")
    run_cierre.bold = True
    sin_espacios(p_cierre)

    # 🔹 FIRMA
    p_firma = doc.add_paragraph(usuario)

    # 🔹 GUARDAR EN BUFFER
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"informe_{correspondenciaElaborada.cite}.docx"
    return buffer, filename

from io import BytesIO
from django.utils.html import strip_tags
import html
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from ..utils.fechas import formatear_fecha_es
from ..utils.formato import sin_espacios, agregar_linea_divisora
from ..word.base import obtener_fecha

def generar_comunicado_word(correspondenciaElaborada):
  
    doc = Document()
    # 🔹 TITULO
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("COMUNICADO")
    run_titulo.bold = True
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_titulo.bold = True
    p_titulo.runs[0].font.size = Pt(16)
    sin_espacios(p_titulo)


    # 🔹 CITE
    p_cite = doc.add_paragraph(correspondenciaElaborada.cite)
    p_cite.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_cite.bold = True
    p_cite.runs[0].font.size = Pt(14)
  

    if hasattr(correspondenciaElaborada, "descripcion_desarrollo"):
        desarrollo = strip_tags(correspondenciaElaborada.descripcion_desarrollo or "")
        p_desarrollo = doc.add_paragraph()
        run_des = p_desarrollo.add_run(f"\n{desarrollo}")
        p_desarrollo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph()  # espacio antes de cierre

    # FECHA
    fecha = obtener_fecha(correspondenciaElaborada.fecha_envio)
    fecha_envio_str = formatear_fecha_es(fecha)
    p_fecha = doc.add_paragraph(f"La Paz, {fecha_envio_str}")
    sin_espacios(p_fecha)
    p_fecha.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_fecha.runs[0].bold = True

    # 🔹 SALUDO DE CIERRE
    p_cierre = doc.add_paragraph()
    run_cierre = p_cierre.add_run("P´ COMITÉ EJECUTIVO")
    p_cierre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_cierre.bold = True
    sin_espacios(p_cierre)

   

    # 🔹 GUARDAR EN BUFFER
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"informe_{correspondenciaElaborada.cite}.docx"
    return buffer, filename
